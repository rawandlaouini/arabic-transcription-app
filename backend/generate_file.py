import sys
import re
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import A4
from bidi.algorithm import get_display
import arabic_reshaper
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import os

# Create downloads directory if it doesn't exist
if not os.path.exists("downloads"):
    os.makedirs("downloads")

# Load Arabic font
pdfmetrics.registerFont(TTFont("Noto", "NotoNaskhArabic-Regular.ttf"))

def generate_file(text, output_format):
    cleaned_text = re.sub(r'\s+', ' ', text.strip())
    if not cleaned_text:
        raise ValueError("Text is empty")

    if output_format == "PDF":
        output_file = os.path.join("downloads", f"audio_arabe_{os.urandom(8).hex()}.pdf")
        pdf = canvas.Canvas(output_file, pagesize=A4)
        pdf.setFont("Noto", 14)
        page_width, page_height = A4
        margin = 50
        x = page_width - margin
        y = page_height - margin
        line_height = 25
        max_width = page_width - 2 * margin
        reshaped_text = arabic_reshaper.reshape(cleaned_text)
        bidi_text = get_display(reshaped_text)
        words = bidi_text.split()
        lines = []
        current_line = []

        for word in words:
            if not current_line:
                test_line = word
            else:
                test_line = " ".join(current_line + [word])
            if pdf.stringWidth(test_line, "Noto", 14) > max_width:
                lines.append(" ".join(current_line))
                current_line = [word]
            else:
                current_line.append(word)

        if current_line:
            lines.append(" ".join(current_line))
        lines = lines[::-1]

        for line in lines:
            if line:
                pdf.drawRightString(x, y, line)
                y -= line_height
                if y < margin:
                    pdf.showPage()
                    pdf.setFont("Noto", 14)
                    y = page_height - margin

        pdf.save()
        return output_file

    else:  # Word format
        output_file = os.path.join("downloads", f"audio_arabe_{os.urandom(8).hex()}.docx")
        doc = Document()
        section = doc.sections[0]
        sectPr = section._sectPr
        bidi = OxmlElement("w:bidi")
        sectPr.append(bidi)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run(cleaned_text)
        run.font.name = "Noto Naskh Arabic"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Naskh Arabic")
        doc.save(output_file)
        return output_file

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python generate_file.py <text> <format>", file=sys.stderr)
        sys.exit(1)

    text = sys.argv[1]
    output_format = sys.argv[2]
    try:
        output_file = generate_file(text, output_format)
        print(output_file)
    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)